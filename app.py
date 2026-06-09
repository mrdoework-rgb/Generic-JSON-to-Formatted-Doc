import streamlit as st
import json
import re
from docx import Document
import io

# --- 1. CHEMISTRY REGEX FUNCTION ---
def append_science_text(paragraph, raw_text):
    """
    Splits text by <sub> tags and applies Microsoft Word subscript formatting.
    """
    # Split text but keep the tags in the resulting list
    parts = re.split(r'(<sub>.*?</sub>)', raw_text)
    
    for part in parts:
        if part.startswith('<sub>') and part.endswith('</sub>'):
            # Strip tags and add as subscript run
            subscript_content = part.replace('<sub>', '').replace('</sub>', '')
            run = paragraph.add_run(subscript_content)
            run.font.subscript = True
        else:
            # Add normal text
            paragraph.add_run(part)

# --- 2. CORE DOCUMENT COMPILER ---
def generate_worksheet(json_data, template_path):
    """
    Loads the template, finds anchor tags, injects data, and returns a BytesIO object.
    """
    doc = Document(template_path)
    
    # Extract data from JSON
    title = json_data.get("worksheet_title", "Science Worksheet")
    body_content = json_data.get("body_content", [])
    summary_table = json_data.get("summary_table", {})
    questions = json_data.get("numbered_questions", [])

    # Iterate through all paragraphs in the document
    for p in doc.paragraphs:
        
        # --- A. TITLE ---
        if "[[WORKSHEET_TITLE]]" in p.text:
            p.text = "" # Clear the anchor
            append_science_text(p, title)
            
        # --- B. BODY CONTENT ---
        elif "[[BODY_CONTENT_HERE]]" in p.text:
            style_name = p.style.name
            for text_block in body_content:
                new_p = p.insert_paragraph_before('')
                new_p.style = style_name
                append_science_text(new_p, text_block)
            
            # Delete the anchor paragraph via XML
            p._element.getparent().remove(p._element)
            
        # --- C. SUMMARY TABLE ---
        elif "[[SUMMARY_TABLE_HERE]]" in p.text:
            rows_data = summary_table.get("rows", [])
            
            if rows_data: # Only build if there is data
                col_count = summary_table.get("column_count", len(rows_data[0]))
                row_count = len(rows_data)
                
                # Word creates tables at the end of the doc by default. 
                # We create it, style it, then move its XML directly after our anchor.
                table = doc.add_table(rows=row_count, cols=col_count)
                try:
                    table.style = 'Science_Table_Style'
                except KeyError:
                    pass # Fallback to default if the style isn't in the template yet
                
                # Move table element right after the anchor paragraph
                p._p.addnext(table._tbl)
                
                # Populate cells and apply chemistry regex
                for r_idx, row_data in enumerate(rows_data):
                    for c_idx, cell_text in enumerate(row_data):
                        cell_paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                        cell_paragraph.text = "" # Clear default empty run
                        append_science_text(cell_paragraph, str(cell_text))
            
            # Delete the anchor paragraph
            p._element.getparent().remove(p._element)
            
        # --- D. QUESTIONS ---
        elif "[[QUESTIONS_HERE]]" in p.text:
            # We assume you've set this anchor to 'List Number' in Word
            style_name = p.style.name 
            for q_text in questions:
                new_p = p.insert_paragraph_before('')
                new_p.style = style_name
                append_science_text(new_p, q_text)
            
            # Delete the anchor paragraph
            p._element.getparent().remove(p._element)

    # Save to an in-memory buffer (Streamlit cloud requires this)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, title

# --- 3. STREAMLIT UI ---
st.set_page_config(page_title="Science Worksheet Generator", page_icon="🔬")

st.title("🔬 Science Worksheet Generator")
st.write("Paste the JSON output from Gemini below to instantly format your printable worksheet.")

# User inputs the JSON
json_input = st.text_area("Paste Gemini JSON here:", height=300)

if st.button("Generate Worksheet", type="primary"):
    if not json_input.strip():
        st.error("Please paste the JSON payload first.")
    else:
        try:
            # Parse the text into a Python Dictionary
            json_data = json.loads(json_input)
            
            # Call the compiler using your template file
            # Make sure 'Generic Template.docx' is in the same GitHub folder as this app.py
            word_buffer, doc_title = generate_worksheet(json_data, "Generic Template.docx")
            
            st.success("Worksheet generated successfully!")
            
            # Provide the download button
            st.download_button(
                label="📥 Download Word Document",
                data=word_buffer,
                file_name=f"{doc_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except json.JSONDecodeError:
            st.error("❌ Invalid JSON. Please check that Gemini output exactly matches the required structure and hasn't cut off.")
        except Exception as e:
            st.error(f"❌ An error occurred while building the document: {e}")
